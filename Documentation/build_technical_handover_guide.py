from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Guide_technique_passation_ATLAS.docx"
IMG_DIR = ROOT / "Images"
QA_DIR = ROOT / "qa_technical_guide"
ARCH_IMG = IMG_DIR / "architecture-atlas.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
MID_GRAY = "667085"
LIGHT_BORDER = "CCD5E0"
WHITE = "FFFFFF"
INK = "202124"
GREEN = "2E6B45"
AMBER = "8A5A00"
RED = "9B1C1C"
MONO_FILL = "F6F8FA"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=INK, bold=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, BLUE, True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, BLUE, True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, "Calibri", 12, DARK_BLUE, True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        set_style_font(style, "Calibri", 11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=9, color=MID_GRAY)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ATLAS  |  GUIDE TECHNIQUE DE PASSATION")
    set_run_font(r, size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer
    p = footer.paragraphs[0]
    p.add_run("Version du 13 août 2026")
    set_run_font(p.runs[0], size=8.5, color=MID_GRAY)
    p.add_run("  |  ")
    set_run_font(p.runs[1], size=8.5, color=MID_GRAY)
    add_page_number(p)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=30, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(30)
        r = p.add_run(subtitle)
        set_run_font(r, size=15, color=DARK_BLUE)


def add_paragraph(doc, text="", bold_prefix=None, italic=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), MONO_FILL)
    p_pr.append(shd)
    r = p.add_run(text)
    set_run_font(r, name="Consolas", size=9, color="1F2937")
    return p


def add_callout(doc, label, text, kind="info"):
    fill = PALE_BLUE
    accent = BLUE
    if kind == "warning":
        fill, accent = "FFF4DB", AMBER
    elif kind == "risk":
        fill, accent = "FDECEC", RED
    elif kind == "success":
        fill, accent = "EAF5EE", GREEN
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), accent)
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(label.upper() + "  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=INK)


def add_table(doc, headers, rows, widths=None, font_size=9.3):
    if widths is None:
        widths = [int(9360 / len(headers))] * len(headers)
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        set_cell_shading(hdr.cells[idx], PALE_BLUE)
        p = hdr.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.2, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], "FAFBFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            is_path = ("/" in str(value) or "\\" in str(value)) and len(str(value)) < 160
            set_run_font(r, name="Consolas" if is_path else "Calibri", size=8.5 if is_path else font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_image(doc, path, caption, width=6.2):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MID_GRAY)


def page_break(doc):
    doc.add_page_break()


def create_architecture_image():
    width, height = 1400, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_candidates = [
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/seguisb.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf"),
    ]
    regular_path = next((p for p in font_candidates if p.exists()), None)
    bold_path = next((p for p in bold_candidates if p.exists()), regular_path)
    regular = ImageFont.truetype(str(regular_path), 27) if regular_path else ImageFont.load_default()
    small = ImageFont.truetype(str(regular_path), 21) if regular_path else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 31) if bold_path else ImageFont.load_default()
    title = ImageFont.truetype(str(bold_path), 39) if bold_path else ImageFont.load_default()

    draw.text((55, 30), "Architecture fonctionnelle d’ATLAS", font=title, fill="#17365D")

    boxes = [
        (50, 130, 370, 265, "MadCap Flare", "IHost • ruban • topic actif"),
        (540, 110, 860, 285, "MyFlarePlugin", "3 commandes du ruban\n+ contexte du projet"),
        (1030, 130, 1350, 265, "Formulaires", "saisie • validation • résumé"),
        (40, 450, 420, 625, "Doc Generator", "AtlasDocGenerationService\n+ Topics • TOC • Target"),
        (510, 450, 890, 625, "Checklist", "ChecklistGeneratorService\n+ H1 • snippets • TOC"),
        (980, 450, 1360, 625, "Finalisation AIT", "AitWorkflowService\n+ Cleanup • ressources • target"),
    ]
    for x1, y1, x2, y2, heading, body in boxes:
        fill = "#E8EEF5" if y1 < 300 else "#F2F4F7"
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill=fill, outline="#2E74B5", width=3)
        hb = draw.textbbox((0, 0), heading, font=bold)
        draw.text(((x1+x2-(hb[2]-hb[0]))/2, y1+22), heading, font=bold, fill="#17365D")
        for idx, line in enumerate(body.split("\n")):
            bb = draw.textbbox((0, 0), line, font=small)
            draw.text(((x1+x2-(bb[2]-bb[0]))/2, y1+75+idx*34), line, font=small, fill="#202124")

    def arrow(a, b):
        draw.line((a, b), fill="#2E74B5", width=6)
        x, y = b
        draw.polygon([(x, y), (x-18, y-10), (x-18, y+10)], fill="#2E74B5")

    arrow((370, 198), (540, 198))
    arrow((860, 198), (1030, 198))
    for end_x in (240, 700, 1160):
        draw.line((700, 285, 700, 365), fill="#2E74B5", width=6)
        draw.line((700, 365, end_x, 365), fill="#2E74B5", width=6)
        draw.line((end_x, 365, end_x, 450), fill="#2E74B5", width=6)
        draw.polygon([(end_x, 450), (end_x-10, 432), (end_x+10, 432)], fill="#2E74B5")

    IMG_DIR.mkdir(parents=True, exist_ok=True)
    image.save(ARCH_IMG)


def add_cover(doc):
    add_title(doc, "ATLAS", "Guide technique du code et dossier de passation")
    add_callout(
        doc,
        "Objectif",
        "Permettre à une développeuse qui découvre le projet de localiser rapidement le bon fichier, comprendre le flux d’exécution et modifier le plugin sans fragiliser les projets MadCap Flare.",
        "info",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    for line, size, bold in [
        ("Nom officiel du plugin : ATLAS", 13, True),
        ("Assembly historique conservé : ATLASDocGenerator.dll", 10.5, False),
        ("Cible : MadCap Flare 21 • .NET Framework 4.8", 10.5, False),
        ("État documenté : 13 août 2026", 10.5, False),
    ]:
        r = p.add_run(line + "\n")
        set_run_font(r, size=size, bold=bold, color=NAVY if bold else MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("Document de maintenance interne")
    set_run_font(r, size=10, italic=True, color=MID_GRAY)


def add_orientation(doc):
    doc.add_heading("1. Lire le projet en 10 minutes", level=1)
    add_paragraph(doc, "ATLAS est une DLL chargée par MadCap Flare. Le ruban ne contient que les points d’entrée ; la logique métier se trouve dans des services indépendants qui lisent et écrivent les fichiers XML de Flare.")
    add_image(doc, ARCH_IMG, "Figure 1 — Séparation entre intégration Flare, formulaires et services métier", 6.55)
    doc.add_heading("Les trois fonctions visibles", level=2)
    add_table(doc, ["Bouton", "Point d’entrée", "Service principal", "Résultat"], [
        ("Doc Generator", "OpenDocGeneratorPopup", "AtlasDocGenerationService", "Dossier documentaire + topics + TOC + target"),
        ("Generate Checklist", "GenerateChecklist", "ChecklistGeneratorService", "Topic checklist ; éventuellement nouvelle TOC et target"),
        ("Finaliser import AIT", "OpenAitWorkflowPopup", "AitWorkflowService", "Ressources, cleanup, TOC, contrôle/réparation target, rapport"),
    ], [1600, 2150, 2350, 3260], 8.8)
    add_callout(doc, "Règle de nommage", "Le produit s’appelle « ATLAS ». Les noms techniques historiques `ATLASDocGenerator`, `ATLASDocGenerator.dll` et les namespaces C# restent inchangés pour éviter une migration risquée.", "warning")
    doc.add_heading("Chemin de lecture conseillé", level=2)
    for item in [
        "Commencer par `MyFlarePlugin.cs` pour voir comment Flare appelle chaque fonction.",
        "Ouvrir ensuite le formulaire associé pour comprendre les champs et validations.",
        "Suivre l’objet Request/Options transmis au service métier.",
        "Lire le service orchestrateur, puis uniquement les sous-services concernés.",
        "Terminer par le fichier de tests portant le même nom fonctionnel.",
    ]:
        add_number(doc, item)


def add_repository_map(doc):
    doc.add_heading("2. Cartographie du dépôt", level=1)
    add_code(doc, "Plugin/\n├─ ATLASDocGenerator/          # DLL chargée par Flare\n├─ ATLASDocGenerator.Tests/    # tests MSTest .NET Framework 4.8\n├─ Documentation/              # guides et scripts de génération\n├─ Kit_Test_Flare/             # recette manuelle reproductible\n└─ ATLASDocGenerator.slnx")
    add_table(doc, ["Zone", "Rôle", "À modifier quand…"], [
        ("MyFlarePlugin.cs", "Entrée IPlugin, onglet ATLAS et boutons du ruban", "on ajoute/renomme un bouton ou change le contexte Flare"),
        ("Forms/", "WinForms : champs, choix, messages, validation immédiate", "on modifie l’interface ou les données demandées"),
        ("Models/", "Objets d’échange Request/Options/Report", "une nouvelle information doit circuler entre UI et service"),
        ("Services/", "Logique métier et orchestration", "on change la génération ou les transformations"),
        ("Services/AitCleanup/", "Transformations XML issues d’Author-it", "on ajoute ou corrige une règle de cleanup"),
        ("Services/AitImportFinalizer/", "Ressources, TOC, target, profils de documents", "on adapte une finalisation AIT"),
        ("Templates/", "Package de ressources copié dans les projets Flare", "CSS, snippets, images ou layouts évoluent"),
        ("ATLASDocGenerator.Tests/", "Tests de non-régression", "chaque changement métier doit être couvert"),
    ], [2200, 3300, 3860], 8.9)
    doc.add_heading("Fichiers à ne pas renommer sans migration", level=2)
    add_bullet(doc, "`ATLASDocGenerator.csproj`, l’assembly `ATLASDocGenerator.dll` et le namespace `ATLASDocGenerator.*`.")
    add_bullet(doc, "Les noms MadCap attendus par les projets : `General.flvar`, extensions `.fltoc`, `.fltar`, `.flpgl`, `.flsnp`.")
    add_bullet(doc, "Les chemins communs `Content/Resources/...` et `Project/...` utilisés dans les targets et TOC.")
    add_callout(doc, "Pourquoi", "Le nom commercial et le nom du binaire peuvent être différents. Le binaire est déjà déployé et référencé ; le renommer n’apporte aucun bénéfice fonctionnel immédiat.", "info")


def add_entry_ribbon(doc):
    doc.add_heading("3. Intégration MadCap Flare et ruban", level=1)
    add_code(doc, "ATLASDocGenerator/MyFlarePlugin.cs")
    add_paragraph(doc, "Cette classe implémente `IPlugin`. Flare appelle `Initialize`, puis `Execute`. `Execute` récupère les contextes d’édition et de navigation, puis crée l’onglet ATLAS.")
    add_table(doc, ["Méthode", "Responsabilité", "Risque si modifiée"], [
        ("GetName", "Nom affiché du plugin : ATLAS", "faible ; garder le nom officiel"),
        ("Initialize", "Mémorise `IHost`", "fort ; aucun accès UI avant l’initialisation"),
        ("Execute", "Récupère les contextes et crée le ruban", "fort ; une exception empêche l’activation"),
        ("CreateAtlasRibbon", "Déclare onglet, groupes, boutons et commandes", "moyen ; vérifier les doublons au rechargement"),
        ("Open…Popup / GenerateChecklist", "Résout le projet actif et ouvre le formulaire", "moyen ; conserver l’exigence d’un topic actif"),
        ("Stop", "Libère l’hôte et désactive le plugin", "fort ; ne pas supprimer le `Dispose`"),
    ], [2200, 3800, 3360], 9)
    doc.add_heading("Résolution du projet actif", level=2)
    add_code(doc, "ATLASDocGenerator/Services/FlareProjectContextService.cs")
    add_paragraph(doc, "Le service part du document actif et remonte l’arborescence jusqu’à trouver une racine Flare cohérente (`Content`, `Project` et/ou `.flprj`). Il charge aussi les dispositifs depuis `General.flvar` pour le Doc Generator.")
    add_callout(doc, "Recette d’ajout d’un quatrième bouton", "Ajouter la commande dans `CreateAtlasRibbon`, créer une méthode d’ouverture, un formulaire dans `Forms/`, un modèle dans `Models/` si nécessaire, puis déléguer toute logique métier à un nouveau service testable.", "success")


def add_doc_generator(doc):
    doc.add_heading("4. Fonction Doc Generator", level=1)
    add_image(doc, IMG_DIR / "doc-generator.png", "Figure 2 — Formulaire de génération documentaire", 5.8)
    add_paragraph(doc, "Le Doc Generator crée un package autonome dans le projet Flare déjà ouvert. Il ne demande plus de choisir manuellement la racine du projet.")
    add_table(doc, ["Étape", "Fichier", "Ce qu’il fait"], [
        ("Interface", "Forms/DocGeneratorForm.cs", "Affiche PS/Notice, titre, référence, dispositif, gamme et titre complet ; valide la saisie"),
        ("Données", "Models/DocGeneratorRequest.cs", "Transporte les champs vers la couche métier"),
        ("Orchestration", "Services/AtlasDocGenerationService.cs", "Valide toutes les dépendances, crée le dossier, pilote topics/TOC/target et rollback"),
        ("Noms", "Services/FileNameSanitizer.cs", "Normalise référence et titre pour les chemins Windows/Flare"),
        ("Topics", "Services/TopicDuplicator.cs", "Définit les listes PS/Notice et copie les modèles"),
        ("Post-traitement", "Services/TopicPostProcessor.cs", "Retire les conditions et répare les chemins de ressources"),
        ("TOC", "Services/TocDuplicator.cs", "Charge la TOC locale ou embarquée, conserve les niveaux et remplace les liens"),
        ("Target", "Services/TargetDuplicator.cs", "Charge la target locale ou embarquée, relie TOC/CSS/layout et variables"),
        ("Fallback", "Services/EmbeddedDocGeneratorTemplates.cs", "Contient les TOC PS/Notice et la target de secours en Base64"),
    ], [1200, 3100, 5060], 8.5)
    doc.add_heading("Ordre de sécurisation", level=2)
    for item in [
        "Valider la demande et calculer tous les chemins cibles.",
        "Vérifier les modèles, la TOC, la target, la CSS, le layout et `General.flvar` avant la première écriture.",
        "Créer le dossier documentaire et les topics.",
        "Créer la TOC puis la target.",
        "En cas d’erreur, supprimer uniquement les artefacts créés pendant cette exécution.",
    ]:
        add_number(doc, item)
    add_callout(doc, "Point important", "Les liens hérités de `Content/Template_tech` dans une TOC embarquée sont des clés de remplacement, pas une obligation de conserver cet ancien dossier. Les sources actuelles sont recherchées sous `Content/Resources/Commun Stago/Topics_Tech` avec compatibilité legacy.", "warning")


def add_doc_generator_changes(doc):
    doc.add_heading("5. Modifier le Doc Generator", level=1)
    add_table(doc, ["Besoin", "Fichier principal", "Autres fichiers / test"], [
        ("Ajouter un type de document", "Forms/DocGeneratorForm.cs", "TopicDuplicator.cs, TocDuplicator.cs, tests DocGeneratorStabilityTests.cs"),
        ("Changer les topics d’un PS", "Services/TopicDuplicator.cs → GetPsRules", "Templates/Commun Stago/Topics_Tech + carte des liens TOC"),
        ("Changer les topics d’une Notice", "Services/TopicDuplicator.cs → GetNoticeRules", "modèle `Notice.fltoc` embarqué/local"),
        ("Ajouter une gamme/CSS", "Forms/DocGeneratorForm.cs", "TargetDuplicator.GetStylesheetPath + fichier Templates/Stylesheets"),
        ("Ajouter un dispositif", "Project/VariableSets/General.flvar", "FlareProjectContextService.LoadDeviceNames"),
        ("Modifier la hiérarchie de TOC", "modèle `.fltoc` source", "TocDuplicator ne doit pas aplatir les `TocEntry` imbriqués"),
        ("Modifier les variables de target", "Services/TargetDuplicator.cs", "DocGeneratorStabilityTests.cs"),
        ("Modifier les modèles de secours", "Services/EmbeddedDocGeneratorTemplates.cs", "recréer les chaînes Base64 et tester sans TOC/target locales"),
    ], [2700, 3250, 3410], 8.5)
    doc.add_heading("Procédure : ajouter un nouveau type", level=2)
    for text in [
        "Ajouter le libellé actif dans `DocGeneratorForm`; ne pas laisser « à venir ».",
        "Ajouter les règles de topics dans `TopicDuplicator.GetRules`.",
        "Fournir une TOC modèle locale et, si l’autonomie est requise, une TOC embarquée.",
        "Vérifier que `TocDuplicator.BuildLinkReplacementMap` connaît tous les liens sources.",
        "Adapter la target si le type exige une CSS, un layout ou des variables différents.",
        "Créer un test complet qui vérifie le nombre de topics, les liens de TOC, la target et le rollback.",
    ]:
        add_number(doc, text)
    doc.add_heading("Pièges connus", level=2)
    add_bullet(doc, "Une entrée ajoutée au formulaire sans règle métier provoque une génération incomplète ou une exception tardive.")
    add_bullet(doc, "Une TOC embarquée peut contenir des chemins legacy : modifier simultanément la carte de remplacement.")
    add_bullet(doc, "Les fichiers `Content` du `.csproj` ne sont pas automatiquement embarqués dans la DLL ; ils sont déployés dans `Templates` à côté du binaire.")
    add_bullet(doc, "Ne pas retirer la validation préalable ni le rollback : ils évitent les demi-documents.")


def add_checklist(doc):
    doc.add_heading("6. Fonction Checklist Generator", level=1)
    add_image(doc, IMG_DIR / "checklist-generator.png", "Figure 3 — Sélection du document et du mode de checklist", 5.7)
    add_table(doc, ["Fichier", "Responsabilité"], [
        ("Forms/ChecklistGeneratorForm.cs", "Charge toutes les targets valides, choisit le mode, demande la nouvelle référence si nécessaire"),
        ("Models/ChecklistGenerationRequest.cs", "Décrit le projet, la target source, le mode et la nouvelle référence"),
        ("Services/Checklist/ChecklistGeneratorService.cs", "Résout target→TOC, parcourt tous les topics, extrait les titres, écrit snippets/topic/TOC/target"),
    ], [3450, 5910], 9)
    doc.add_heading("Règles de collecte", level=2)
    add_bullet(doc, "La source documentaire est la TOC reliée à la target choisie, pas seulement le topic actif.")
    add_bullet(doc, "Tous les `h1` sont collectés dans l’ordre de la TOC, sauf les `h1.no_num` et variantes non numérotées.")
    add_bullet(doc, "Dans un topic Prérequis, les paragraphes `p.ss_section` deviennent des actions de checklist.")
    add_bullet(doc, "Le Sommaire est exclu de la checklist autonome.")
    add_bullet(doc, "Chaque action reçoit un identifiant `checklist-step-001`, utilisé par les entrées enfant de la TOC autonome.")
    add_bullet(doc, "Le tableau QIQO est inséré via `QIQO_table.flsnp`; introduction et titre sont des `snippetText` modifiables.")
    doc.add_heading("Deux modes", level=2)
    add_table(doc, ["Mode", "Écritures"], [
        ("Ajouter au document actuel", "Crée/remplace le topic Checklist dans le dossier principal et ajoute une entrée en fin de TOC"),
        ("Nouveau document", "Crée uniquement le topic checklist, une TOC avec les étapes numérotées et une target dupliquée ; aucun Sommaire visible"),
    ], [2500, 6860], 9)


def add_checklist_changes(doc):
    doc.add_heading("7. Modifier la Checklist", level=1)
    add_table(doc, ["Besoin", "Méthode / fichier", "Test à adapter"], [
        ("Changer les titres pris en compte", "CollectSections / IsNonNumberedH1", "ExcludesNonNumberedH1…"),
        ("Changer les prérequis", "IsPrerequisiteTopic + collecte `p.ss_section`", "…IncludesPrerequisites"),
        ("Changer le texte d’introduction", "EnsureChecklistSnippets → intro_checklist.flsnp", "test du contenu et du type snippetText"),
        ("Changer le tableau QIQO", "Ressource QIQO_content/QIQO_table.flsnp", "test d’existence et référence snippetBlock"),
        ("Changer la structure du topic", "WriteChecklistTopic", "tests XML du topic généré"),
        ("Changer la TOC autonome", "CreateStandaloneChecklistToc", "Generate_NewDocument_ExcludesSummary…"),
        ("Changer la duplication target", "branche CreateNewDocument dans Generate", "…OnlyChangesDocumentReferenceAndToc"),
        ("Changer la découverte des documents", "GetAvailableTargets", "GetAvailableTargets_FindsTargetsInSubfolders"),
    ], [2550, 3850, 2960], 8.4)
    add_callout(doc, "Invariant", "En mode nouveau document, ne jamais recopier les topics du document source. Le package doit contenir seulement la checklist, sa TOC et sa target.", "risk")
    doc.add_heading("Idempotence", level=2)
    add_paragraph(doc, "Une seconde exécution doit remplacer proprement la checklist existante, conserver une seule entrée de TOC et ne pas multiplier les sauvegardes initiales. Toute évolution doit préserver ce comportement.")
    add_code(doc, "Sauvegardes Checklist : <fichier>.before-checklist.bak")


def add_ait(doc):
    doc.add_heading("8. Fonction Finaliser import AIT", level=1)
    add_image(doc, IMG_DIR / "finaliser-import-ait.png", "Figure 4 — Workflow unifié de finalisation Author-it", 5.75)
    add_paragraph(doc, "Le ruban ouvre uniquement `AitWorkflowForm`. Les anciens formulaires `AitCleanupForm` et `AitImportFinalizerForm` sont encore compilés pour compatibilité, mais ils ne sont pas les points d’entrée utilisateur principaux.")
    add_table(doc, ["Couche", "Fichier", "Responsabilité"], [
        ("Interface", "Forms/AitWorkflowForm.cs", "Détecte TOC/targets, options Cleanup/IHM, construit AitWorkflowOptions"),
        ("Orchestrateur", "Services/AitWorkflowService.cs", "Exécute chaque bloc avec gestion d’erreurs séparée et produit un rapport unifié"),
        ("Profils", "AitImportFinalizer/AitDocumentProfileFactory.cs", "CSS, layout et entrées TOC parasites par type documentaire"),
        ("Ressources", "AitImportFinalizer/ResourceCopyService.cs", "Copie Templates vers le projet, compare et sauvegarde avant remplacement"),
        ("TOC", "AitImportFinalizer/TocCleanerService.cs", "Supprime les entrées configurées sans aplatir la hiérarchie"),
        ("Target", "AitImportFinalizer/TargetConfiguratorService.cs", "Valide sans écrire ; répare seulement sur option"),
        ("Cleanup", "AitCleanup/AitCleanupService.cs", "Scanne topics/snippets puis lance les transformeurs choisis"),
    ], [1500, 3400, 4460], 8.2)
    doc.add_heading("Ordre d’exécution", level=2)
    for text in [
        "Installer ou mettre à jour les ressources ATLAS.",
        "Nettoyer les topics et snippets ; traiter éventuellement les variables IHM.",
        "Nettoyer les entrées parasites de la TOC.",
        "Valider la target sans écriture.",
        "Réparer la target uniquement si l’option est cochée et qu’une différence existe.",
        "Écrire le rapport final unifié.",
    ]:
        add_number(doc, text)


def add_cleanup(doc):
    doc.add_heading("9. Moteur Cleanup : où modifier une transformation", level=1)
    add_table(doc, ["Transformation", "Fichier", "Sortie principale"], [
        ("Actions / résultats", "ActionResultListDetector.cs + ActionResultListTransformer.cs", "liste ordonnée `Action_num` et structure résultat"),
        ("Listes à tirets", "BulletListTransformer.cs", "listes `<ul>/<li>` MadCap propres"),
        ("Callouts", "CalloutTransformer.cs", "blocs callout avec classes et pictogrammes attendus"),
        ("Figures", "FigureTransformer.cs", "bloc `div.a_figure`, image et légende"),
        ("Styles simples", "SimpleStyleCleanupTransformer.cs", "suppression classes parasites, `sub`, `sup`, centrage"),
        ("Périmètre fichiers", "HtmlFileScanner.cs", "topics `.htm/.html` valides dans le dossier sélectionné"),
        ("Orchestration", "AitCleanupService.cs", "ordre des transformeurs et rapport"),
        ("Journal", "CleanupLogService.cs", "log détaillé des fichiers, transformations, avertissements et erreurs"),
    ], [2200, 3800, 3360], 8.5)
    doc.add_heading("Patron obligatoire pour une nouvelle règle", level=2)
    for text in [
        "Détecter strictement la structure Author-it source avant de modifier.",
        "Charger le fichier en XML et valider la sortie avant toute sauvegarde.",
        "Créer une sauvegarde initiale immuable via `FileBackupService`.",
        "Rendre la transformation idempotente : un second passage ne doit rien dégrader ni dupliquer.",
        "Compter les changements et alimenter `CleanupReport`.",
        "Ajouter un test premier passage + second passage + sauvegarde.",
    ]:
        add_number(doc, text)
    add_callout(doc, "Attention", "Les fichiers `.flsnp` doivent être parcourus lorsque la règle concerne les références IHM. Ne pas limiter les recherches aux seuls `.htm`.", "warning")


def add_ihm_resources(doc):
    doc.add_heading("10. Variables IHM et ressources ATLAS", level=1)
    doc.add_heading("Chaîne IHM", level=2)
    add_table(doc, ["Étape", "Fichier", "Rôle"], [
        ("Détection", "FrenchIhmTemplateDetector.cs", "Lit le XML Author-it, trouve les templates français réellement utilisés"),
        ("Génération", "FrenchIhmVariableSetGenerator.cs", "Crée un `.flvar` par template sélectionné et extrait les définitions"),
        ("Remplacement", "IhmVariableReferenceTransformer.cs", "Parcourt topics et snippets, remplace les références de snippet par des variables MadCap"),
        ("Rapport", "CleanupReport.cs / CleanupLogService.cs", "Expose fichiers scannés, modifiés, références remplacées et IDs non associés"),
    ], [1500, 3650, 4210], 8.5)
    add_callout(doc, "État actuel", "La logique de conditions IHM n’est pas implémentée faute de règle métier confirmée. Ne pas l’inférer depuis le XML sans nouvelles instructions fonctionnelles.", "warning")
    doc.add_heading("Package Templates", level=2)
    add_code(doc, "ATLASDocGenerator/Templates/\n├─ Stylesheets/      → Content/Resources/Stylesheets\n├─ PageLayouts/      → Content/Resources/PageLayouts\n├─ Snippets/         → Content/Resources/Snippets\n├─ Images/           → Content/Resources/Images\n├─ VariableSets/     → Project/VariableSets\n└─ Commun Stago/     → Content/Resources/Commun Stago")
    add_paragraph(doc, "`ResourceCopyService` localise `Templates` à côté de la DLL avec `Assembly.GetExecutingAssembly().Location`. Le déploiement doit donc copier la DLL et le dossier `Templates` ensemble.")
    add_bullet(doc, "Toutes les sources sont validées avant la première copie.")
    add_bullet(doc, "Un fichier identique est ignoré ; un fichier différent reçoit une sauvegarde initiale avant remplacement.")
    add_bullet(doc, "Un `General.flvar` existant est toujours préservé car il appartient au contenu commun du projet.")
    add_bullet(doc, "La présence de la CSS et du layout du profil est revalidée après copie.")


def add_safety(doc):
    doc.add_heading("11. Sécurité des fichiers et conventions XML", level=1)
    add_table(doc, ["Mécanisme", "Où", "But"], [
        ("Validation avant écriture", "Doc Generator, ResourceCopyService, TargetConfiguratorService", "éviter les packages partiels"),
        ("Sauvegarde initiale", "FileBackupService.cs", "ne jamais écraser la première version de secours"),
        ("Rollback", "AtlasDocGenerationService.cs", "supprimer uniquement les artefacts générés après une erreur"),
        ("Validation seule", "TargetConfiguratorService.ValidateTarget", "afficher les écarts sans modifier la target"),
        ("Préservation XML", "XDocument + PreserveWhitespace/DisableFormatting", "limiter les reformattages inutiles"),
        ("Comparaison insensible à la casse", "attributs et noms MadCap", "tolérer les variantes de Flare sans casser les namespaces"),
    ], [2200, 3600, 3560], 8.6)
    doc.add_heading("Extensions de sauvegarde rencontrées", level=2)
    add_code(doc, ".bak\n.before-ait-finalizer.bak\n.before-checklist.bak\n(les transformeurs Cleanup utilisent également une sauvegarde initiale dédiée)")
    doc.add_heading("Règles pour modifier du XML MadCap", level=2)
    add_bullet(doc, "Comparer `Name.LocalName` lorsque les namespaces peuvent varier.")
    add_bullet(doc, "Appeler `.ToList()` avant de supprimer des éléments pendant une itération LINQ.")
    add_bullet(doc, "Préserver les fragments `#id` et les chemins `/Content/...` attendus par les TOC.")
    add_bullet(doc, "Ne jamais aplatir une TOC : l’imbrication des `TocEntry` représente les niveaux du PDF.")
    add_bullet(doc, "Valider un XML complet en mémoire avant de créer la sauvegarde et d’écrire.")
    add_callout(doc, "General.flvar", "Le Finalizer ne doit pas remplacer ce fichier lors de la copie des ressources. La mise à jour de variables, lorsqu’elle est demandée, se fait de manière ciblée et le MRef est laissé intact pour un manuel de référence.", "risk")


def add_build_tests(doc):
    doc.add_heading("12. Compiler, tester et déployer", level=1)
    doc.add_heading("Environnement", level=2)
    add_table(doc, ["Élément", "Valeur"], [
        ("Framework plugin", ".NET Framework 4.8"),
        ("API hôte", "B3.PluginAPIKit.dll de MadCap Flare 21"),
        ("Interface", "Windows Forms"),
        ("Tests", "MSTest.Sdk 4.3.2, cible net48"),
        ("Sortie Release", "ATLASDocGenerator/bin/Release/ATLASDocGenerator.dll"),
    ], [2600, 6760], 9)
    doc.add_heading("Commandes de référence", level=2)
    add_code(doc, "dotnet build .\\ATLASDocGenerator\\ATLASDocGenerator.csproj -c Release --no-restore\n\ndotnet test .\\ATLASDocGenerator.Tests\\ATLASDocGenerator.Tests.csproj -c Release --no-restore")
    add_paragraph(doc, "Les sources métier sont liées directement au projet de tests. Cela évite de charger la DLL complète du plugin hors de Flare, où la dépendance B3.PluginAPIKit et les politiques d’entreprise peuvent bloquer le chargement.")
    doc.add_heading("Carte des tests", level=2)
    add_table(doc, ["Zone", "Fichier de tests"], [
        ("Doc Generator et rollback", "DocGeneratorStabilityTests.cs"),
        ("Checklist, targets, H1, prérequis", "ChecklistGeneratorServiceTests.cs"),
        ("Transformations Cleanup", "AitCleanupTransformerSafetyTests.cs + SimpleStyleCleanupTransformerTests.cs"),
        ("IHM", "FrenchIhmTemplateDetectorTests.cs + IhmVariableReferenceTransformerTests.cs"),
        ("Ressources", "ResourceCopyServiceTests.cs"),
        ("TOC, variables, target", "FinalizerFileSafetyTests.cs + TargetConfiguratorServiceTests.cs"),
        ("Contexte Flare", "FlareProjectContextServiceTests.cs"),
        ("Recette globale", "ManualFlareKitTests.cs"),
    ], [3300, 6060], 8.8)


def add_recipes(doc):
    doc.add_heading("13. Recettes de maintenance courantes", level=1)
    recipes = [
        ("Ajouter une option au formulaire", "Ajouter le contrôle dans `Forms/...Form.cs`, la propriété dans `Models/...Options.cs`, la transmettre lors du clic, la consommer dans le service et tester les deux états."),
        ("Ajouter un pictogramme de callout", "Ajouter l’image dans `Templates/Images`, vérifier son inclusion/copie dans le `.csproj`, puis adapter `CalloutTransformer.cs` et un test avec chemin relatif."),
        ("Changer une CSS ou un layout AIT", "Modifier le fichier sous `Templates`, puis le chemin du profil dans `AitDocumentProfileFactory.cs`; tester copie + validation/réparation target."),
        ("Ajouter une entrée parasite de TOC", "Ajouter le motif dans `TocEntriesToRemove` du profil concerné et ajouter un cas hiérarchique dans `FinalizerFileSafetyTests.cs`."),
        ("Modifier une variable General", "Adapter `VariableSetUpdaterService.cs`; préserver les structures avec/sans `VariableDefinition` et les règles spécifiques au Manuel de référence."),
        ("Changer la source des dispositifs", "Adapter `FlareProjectContextService.LoadDeviceNames`; conserver `Multi` et `Autre` si la saisie libre reste nécessaire."),
        ("Ajouter un nouveau fichier C#", "Créer le fichier, l’ajouter explicitement au groupe `<Compile Include=...>` de l’ancien `.csproj`, puis compiler Release."),
        ("Ajouter une ressource", "Créer le fichier sous `Templates`, l’ajouter au `.csproj` avec le bon comportement de copie et vérifier le package Release."),
    ]
    for title, body in recipes:
        doc.add_heading(title, level=2)
        add_paragraph(doc, body)
    add_callout(doc, "Avant commit", "Compiler, lancer tous les tests, exécuter `git diff --check`, vérifier le kit Flare si la modification touche l’interface ou les ressources, puis fournir la nouvelle DLL avec son dossier Templates.", "success")


def add_known_limits(doc):
    doc.add_heading("14. Limites connues et décisions à préserver", level=1)
    add_table(doc, ["Sujet", "État / décision"], [
        ("Nom du plugin", "Nom officiel : ATLAS. Les identifiants techniques historiques restent inchangés."),
        ("Liste améliorations/corrections", "Visible comme « à venir » dans le Doc Generator ; aucun modèle métier final n’est fourni."),
        ("Conditions IHM", "En attente d’instructions fiables ; ne pas inventer de logique de filtrage."),
        ("Formulaires AIT historiques", "Toujours présents dans le projet mais le ruban utilise le workflow unifié `AitWorkflowForm`."),
        ("Ressources", "Le dossier `Templates` reste nécessaire à côté de la DLL pour la finalisation AIT."),
        ("TOC/target Doc Generator", "Des modèles de secours sont embarqués ; les modèles locaux du projet restent prioritaires."),
        ("General.flvar", "Toujours préservé pendant la copie des ressources ; mises à jour ciblées uniquement."),
        ("Target Finalizer", "Validation par défaut ; réparation uniquement sur choix explicite de l’utilisateur."),
    ], [2700, 6660], 8.7)
    doc.add_heading("Checklist de passation", level=2)
    for text in [
        "Ouvrir la solution et compiler la configuration Release.",
        "Exécuter la suite complète des tests et lire au moins un test par fonction.",
        "Installer la DLL et le dossier Templates dans un environnement Flare de test.",
        "Tester les trois boutons depuis un topic actif.",
        "Conserver un projet de recette sacrifiable ; ne pas tester les transformations sur un projet sans sauvegarde.",
        "Documenter toute nouvelle règle métier dans ce guide et dans un test nommé explicitement.",
        "Utiliser « ATLAS » dans tous les futurs écrans et documents utilisateur.",
    ]:
        add_bullet(doc, "☐ " + text)
    add_callout(doc, "Principe directeur", "Les formulaires collectent et valident ; les modèles transportent ; les services transforment ; les tests protègent. Si une nouvelle fonctionnalité mélange ces rôles, la découper avant de l’étendre.", "info")


def add_appendix(doc):
    doc.add_heading("Annexe A — Index besoin → fichier", level=1)
    rows = [
        ("Onglet et boutons du ruban", "ATLASDocGenerator/MyFlarePlugin.cs"),
        ("Détection du projet ouvert", "Services/FlareProjectContextService.cs"),
        ("Sélecteur de dossiers moderne", "Services/ModernFolderPicker.cs"),
        ("Champs Doc Generator", "Forms/DocGeneratorForm.cs"),
        ("Création package documentaire", "Services/AtlasDocGenerationService.cs"),
        ("Liste des topics PS/Notice", "Services/TopicDuplicator.cs"),
        ("Chemins des ressources des topics", "Services/TopicPostProcessor.cs"),
        ("Liens et niveaux TOC générée", "Services/TocDuplicator.cs"),
        ("CSS/layout/variables target générée", "Services/TargetDuplicator.cs"),
        ("TOC/target de secours", "Services/EmbeddedDocGeneratorTemplates.cs"),
        ("Formulaire Checklist", "Forms/ChecklistGeneratorForm.cs"),
        ("Extraction H1/prérequis et génération", "Services/Checklist/ChecklistGeneratorService.cs"),
        ("Formulaire finalisation AIT", "Forms/AitWorkflowForm.cs"),
        ("Orchestration AIT", "Services/AitWorkflowService.cs"),
        ("Profils documentaires AIT", "Services/AitImportFinalizer/AitDocumentProfileFactory.cs"),
        ("Copie ressources", "Services/AitImportFinalizer/ResourceCopyService.cs"),
        ("Nettoyage TOC", "Services/AitImportFinalizer/TocCleanerService.cs"),
        ("Validation/réparation target", "Services/AitImportFinalizer/TargetConfiguratorService.cs"),
        ("Variables General", "Services/AitImportFinalizer/VariableUpdaterService.cs"),
        ("Actions/résultats", "Services/AitCleanup/ActionResultListTransformer.cs"),
        ("Callouts", "Services/AitCleanup/CalloutTransformer.cs"),
        ("Figures", "Services/AitCleanup/FigureTransformer.cs"),
        ("Variables IHM", "Services/AitCleanup/IhmVariables/"),
        ("Sauvegardes", "Services/FileBackupService.cs"),
        ("Ressources déployées", "ATLASDocGenerator/Templates/"),
        ("Suite de tests", "ATLASDocGenerator.Tests/"),
    ]
    add_table(doc, ["Je veux modifier…", "Ouvrir d’abord…"], rows, [3700, 5660], 8.3)
    doc.add_heading("Annexe B — Diagnostic rapide", level=1)
    add_table(doc, ["Symptôme", "Premier contrôle"], [
        ("Le bouton n’apparaît pas", "MyFlarePlugin.Execute/CreateAtlasRibbon, activation du plugin, chargement de la DLL"),
        ("Projet actif introuvable", "Ouvrir un topic puis vérifier FlareProjectContextService"),
        ("Topic modèle introuvable", "Vérifier Resources/Commun Stago/Topics_Tech puis compatibilité Template_tech"),
        ("TOC générée avec ancien lien", "Ajouter la clé dans TocDuplicator.BuildLinkReplacementMap"),
        ("CSS/layout absent après Finalizer", "Vérifier dossier Templates à côté de la DLL et profil documentaire"),
        ("Target non conforme", "Lancer validation, lire Differences, cocher réparation si souhaitée"),
        ("Callout sans picto", "Vérifier Templates/Images, copie vers Content/Resources/Images et chemin du transformeur"),
        ("Checklist vide", "Vérifier target→TOC, liens des topics, h1.no_num et p.ss_section dans Prérequis"),
        ("Références IHM non remplacées", "Vérifier XML source, template sélectionné et parcours des `.flsnp`"),
    ], [3600, 5760], 8.5)


def build_document():
    create_architecture_image()
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    add_cover(doc)
    page_break(doc)
    add_orientation(doc)
    page_break(doc)
    add_repository_map(doc)
    page_break(doc)
    add_entry_ribbon(doc)
    page_break(doc)
    add_doc_generator(doc)
    page_break(doc)
    add_doc_generator_changes(doc)
    page_break(doc)
    add_checklist(doc)
    page_break(doc)
    add_checklist_changes(doc)
    page_break(doc)
    add_ait(doc)
    page_break(doc)
    add_cleanup(doc)
    page_break(doc)
    add_ihm_resources(doc)
    page_break(doc)
    add_safety(doc)
    page_break(doc)
    add_build_tests(doc)
    page_break(doc)
    add_recipes(doc)
    page_break(doc)
    add_known_limits(doc)
    page_break(doc)
    add_appendix(doc)

    doc.core_properties.title = "ATLAS - Guide technique du code et dossier de passation"
    doc.core_properties.subject = "Architecture, maintenance, tests et déploiement du plugin ATLAS"
    doc.core_properties.author = "Équipe ATLAS"
    doc.core_properties.keywords = "ATLAS, MadCap Flare, plugin, passation, maintenance, C#"
    doc.core_properties.comments = "Nom officiel du plugin : ATLAS."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
