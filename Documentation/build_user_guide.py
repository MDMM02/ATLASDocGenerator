from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
IMAGES = ROOT / "Images"
OUTPUT = ROOT / "Guide_utilisation_ATLAS_Doc_Generator.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "202124"
WHITE = "FFFFFF"
GOLD = "A66A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)


def add_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "number" else "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    set_run_font(p.add_run(label + "  "), bold=True, color=accent)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_placeholder(doc, title, instructions):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=10, after=4, line=1.15)
    set_run_font(p.add_run("CAPTURE À AJOUTER\n"), bold=True, color=GOLD, size=11)
    set_run_font(p.add_run(title + "\n"), bold=True, color=NAVY, size=11)
    set_run_font(p.add_run(instructions), italic=True, color=MID_GRAY, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, image_name, caption, width):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    picture = run.add_picture(str(IMAGES / image_name), width=Inches(width))
    picture._inline.docPr.set("title", caption)
    picture._inline.docPr.set("descr", caption)
    c = doc.add_paragraph(caption, style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.05)
        set_run_font(p.add_run(text), bold=True, color=NAVY, size=9.5)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.08)
            set_run_font(p.add_run(value), size=9.3)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        set_run_font(p.add_run(bold_start), bold=True)
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0, line=1.0)
    set_run_font(p.add_run("ATLAS DOC GENERATOR  |  GUIDE UTILISATEUR"), size=8.5, color=MID_GRAY, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.0)
    set_run_font(p.add_run("Usage interne  •  Page "), size=9, color=MID_GRAY)
    add_field(p, "PAGE")


def build_document():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_id = add_numbering(doc, "bullet")
    number_id = add_numbering(doc, "number")

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("GUIDE UTILISATEUR"), size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("ATLAS Doc Generator"), size=30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    set_run_font(p.add_run("Utiliser les trois outils du ruban dans MadCap Flare"), size=15, color=DARK_BLUE)

    add_callout(
        doc,
        "OBJECTIF",
        "Créer un document, générer sa checklist et finaliser un import Author-it avec une procédure guidée et sécurisée.",
        fill=LIGHT_BLUE,
        accent=NAVY,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("Version du guide : 11 août 2026"), size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Plug-in 0.1 • MadCap Flare • Usage interne"), size=10, color=MID_GRAY)

    doc.add_page_break()

    # Orientation
    add_heading(doc, "1. Bien démarrer", 1)
    add_body(doc, "Le plug-in ajoute un onglet ATLAS au ruban de MadCap Flare. Il fonctionne dans le projet auquel appartient le topic actuellement ouvert : le dossier du projet n'est donc pas demandé dans les formulaires.")
    add_callout(doc, "PRÉREQUIS", "Avant de cliquer sur un bouton ATLAS, ouvrir dans l’éditeur un topic appartenant au projet Flare à traiter.")

    add_heading(doc, "1.1 Les trois boutons", 2)
    add_table(
        doc,
        ["Bouton", "Quand l'utiliser", "Résultat principal"],
        [
            ("Doc Generator", "Au démarrage d'un nouveau document.", "Crée le dossier de contenu, les topics, la TOC et la target."),
            ("Generate Checklist", "Quand le document source et sa TOC sont structurés.", "Crée une checklist à partir des étapes du document."),
            ("Finaliser import AIT", "Après un import Author-it dans Flare.", "Installe les ressources, nettoie le contenu et contrôle la TOC et la target."),
        ],
        [2100, 3300, 3960],
    )
    add_placeholder(
        doc,
        "Figure 1 — Onglet ATLAS dans le ruban Flare",
        "Ouvrir un projet dans Flare, sélectionner l'onglet ATLAS et capturer les trois groupes : Documentation, Checklist et Author-it.",
    )

    add_heading(doc, "1.2 Règles générales de sécurité", 2)
    for text in (
        "Travailler sur une copie du projet lors du premier essai d'un import Author-it.",
        "Fermer ou enregistrer les topics concernés avant une transformation massive.",
        "Lire le récapitulatif et les journaux avant de relancer une opération.",
        "Ne pas supprimer les fichiers .bak : ils servent de sauvegarde initiale.",
        "Conserver le dossier Templates à côté de ATLASDocGenerator.dll pour que les ressources puissent être installées.",
    ):
        add_list_item(doc, text, bullet_id)

    # Doc Generator
    add_heading(doc, "2. Doc Generator", 1, page_break=True)
    add_body(doc, "Doc Generator initialise un nouveau package documentaire dans le projet Flare actif. Il valide les modèles avant d'écrire, crée les éléments nécessaires, puis annule les fichiers créés si une étape échoue.")
    add_figure(doc, "doc-generator.png", "Figure 2 — Formulaire Doc Generator", 6.2)

    add_heading(doc, "2.1 Que saisir dans chaque champ", 2)
    add_table(
        doc,
        ["Champ", "Valeur attendue", "Exemple / règle"],
        [
            ("Type de document", "PS ou Notice.", "Liste des améliorations / corrections est visible mais reste indisponible."),
            ("Titre doc abrégé", "Titre court servant au nom du dossier et des fichiers.", "40 caractères maximum. Ne pas utiliser < > : \" / \\ | ? *. Exemple : Installation_module."),
            ("Référence sans indice", "Référence documentaire sans numéro de révision.", "Exemple : PS_STHEMO_001. Elle entre dans les noms générés."),
            ("Dispositif", "Choisir le dispositif dans la liste issue de General.flvar.", "La liste ajoute toujours Multi et Autre."),
            ("Nom dispositif libre", "À remplir seulement pour Multi ou Autre.", "Saisir le libellé exact à afficher dans le document."),
            ("Gamme", "Choisir sthemX ou STA.", "STA utilise Styles_STA.css ; sthemX utilise Styles.css."),
            ("Titre complet", "Titre destiné au document.", "120 caractères maximum. Les accents et espaces sont acceptés."),
        ],
        [1850, 3650, 3860],
    )

    add_heading(doc, "2.2 Procédure", 2)
    for text in (
        "Ouvrir un topic du projet dans lequel le document doit être créé.",
        "Dans le ruban ATLAS, cliquer sur Doc Generator.",
        "Choisir le type documentaire et renseigner tous les champs.",
        "Vérifier particulièrement la référence sans indice et le titre abrégé : ils déterminent les noms de fichiers.",
        "Cliquer sur Générer.",
        "Lire le récapitulatif indiquant le dossier, le nombre de topics, la TOC et la target créés.",
    ):
        add_list_item(doc, text, number_id)

    add_heading(doc, "2.3 Résultat créé", 2)
    add_callout(doc, "NOM DU PACKAGE", "Le dossier, la TOC et la target sont nommés <Référence>_<Titre abrégé> après normalisation des caractères.")
    add_table(
        doc,
        ["Emplacement", "Contenu"],
        [
            ("Content/<Référence>_<Titre abrégé>/", "Topics de titre, historique, objectif, sécurité, matériel, documents, durée/remplacements, prérequis et premier chapitre selon le type."),
            ("Project/TOCs/", "TOC portant le même nom que le package."),
            ("Project/Targets/", "Target PDF portant le même nom et liée à la nouvelle TOC."),
        ],
        [3000, 6360],
    )

    add_heading(doc, "2.4 Contrôles après génération", 2)
    for text in (
        "Actualiser l'arborescence Flare si les nouveaux fichiers ne sont pas immédiatement visibles.",
        "Ouvrir la TOC et vérifier l'ordre des topics.",
        "Ouvrir la target et confirmer la TOC, la CSS et le page layout.",
        "Ouvrir quelques topics et contrôler les images, snippets et liens.",
        "Lancer un build PDF de contrôle.",
    ):
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "EN CAS DE COLLISION", "Le générateur refuse d'écraser un dossier, une TOC ou une target existante. Corriger la référence ou le titre abrégé, ou traiter manuellement l'ancien package.", fill="FFF7E6", accent=GOLD)

    # Checklist
    add_heading(doc, "3. Generate Checklist", 1, page_break=True)
    add_body(doc, "Checklist Generator parcourt toute la TOC du document sélectionné. Il récupère les H1 admissibles et transforme les paragraphes p.ss_section du topic Prérequis en actions de checklist.")
    add_figure(doc, "checklist-generator.png", "Figure 3 — Formulaire Checklist Generator", 6.2)

    add_heading(doc, "3.1 Sélectionner le document", 2)
    add_body(doc, "La liste affiche les targets trouvées dans Project/Targets et ses sous-dossiers. Une target n'est proposée que si son attribut MasterToc pointe vers une TOC valide. Le nom peut être complété par la variable General/DocumentReference lorsqu'elle existe dans la target.")

    add_heading(doc, "3.2 Choisir le mode", 2)
    add_table(
        doc,
        ["Mode", "À utiliser lorsque", "Ce que le plug-in crée"],
        [
            ("Ajouter à la fin de la TOC actuelle", "La checklist appartient au même document.", "Checklist.htm dans le dossier principal et une entrée Checklist à la fin de la TOC existante."),
            ("Créer un nouveau document checklist", "La checklist doit être publiée séparément.", "Un dossier <Document Reference>_checklist, une TOC dédiée et une target dupliquée depuis le document source."),
        ],
        [2600, 2900, 3860],
    )
    add_callout(doc, "DOCUMENT REFERENCE", "Ce champ est obligatoire uniquement pour un nouveau document. Saisir la référence du futur document checklist ; le suffixe _checklist est ajouté automatiquement.")

    add_heading(doc, "3.3 Contenu pris en compte", 2)
    for text in (
        "Tous les H1 des topics référencés par la TOC sélectionnée.",
        "Les H1 portant les classes no_num, non_numerote, non_numéroté ou non_num sont ignorés.",
        "Dans le topic Prérequis, seuls les paragraphes de classe p.ss_section deviennent des actions.",
        "En mode nouveau document, le H1 Sommaire n'est pas ajouté à la checklist.",
        "Une ancienne Checklist.htm n'est pas relue comme source d'actions.",
    ):
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "3.4 Structure produite", 2)
    add_body(doc, "La checklist contient uniquement le titre, la ligne d'introduction et une liste numérotée d'actions. Chaque action reçoit le snippet QIQO_table. Les snippets titre_checklist.flsnp et intro_checklist.flsnp sont créés ou actualisés dans Resources/Commun Stago/QIQO_content.")
    add_callout(doc, "PRÉREQUIS QIQO", "Le fichier Content/Resources/Commun Stago/QIQO_content/QIQO_table.flsnp doit exister. Sans lui, la génération est arrêtée avec un message explicite.")

    add_heading(doc, "3.5 Procédure", 2)
    for text in (
        "Ouvrir un topic du projet et cliquer sur Generate Checklist.",
        "Sélectionner la target correspondant au document source.",
        "Choisir l'ajout à la TOC actuelle ou la création d'un nouveau document.",
        "Pour un nouveau document, renseigner Document Reference.",
        "Cliquer sur Générer puis contrôler le nombre de sections annoncé.",
        "Vérifier Checklist.htm, la TOC et, si elle a été créée, la nouvelle target.",
        "Construire le PDF et vérifier la numérotation des étapes et les tableaux QIQO.",
    ):
        add_list_item(doc, text, number_id)

    # AIT
    add_heading(doc, "4. Finaliser import AIT", 1, page_break=True)
    add_body(doc, "Ce bouton regroupe l'ancien AIT Cleanup et l'ancien AIT Import Finalizer. Il permet de choisir précisément les traitements à appliquer après un import Author-it.")
    add_figure(doc, "finaliser-import-ait.png", "Figure 4 — Formulaire unifié Finaliser import AIT", 5.25)

    add_heading(doc, "4.1 Document et fichiers", 2)
    add_table(
        doc,
        ["Champ", "Sélection attendue"],
        [
            ("Type de document", "Choisir le profil correspondant : Bulletin Technique, Notice utilisateur, Addenda, Manuel de référence / MRef, Document technique ou Document technique multi-instrument."),
            ("TOC", "Choisir la TOC importée à nettoyer et à contrôler. Les sous-dossiers de Project/TOCs sont inclus."),
            ("Target", "Choisir la target du même document, généralement la target PDF. Lors du changement de target, le plug-in tente de sélectionner automatiquement sa MasterToc."),
            ("Dossier à nettoyer", "Choisir le dossier contenant les topics importés. Par défaut, tout le dossier Content du projet actif est proposé."),
        ],
        [2350, 7010],
    )

    add_heading(doc, "4.2 Actions principales", 2)
    add_table(
        doc,
        ["Option", "Défaut", "Effet"],
        [
            ("Installer ou mettre à jour les ressources ATLAS", "Oui", "Copie les page layouts, CSS, snippets, images, VariableSets et ressources Commun Stago livrés à côté de la DLL. General.flvar existant est toujours conservé."),
            ("Nettoyer les topics et snippets", "Oui", "Lance les transformations Cleanup sélectionnées sur le dossier choisi. Le cleanup standard analyse les topics .htm ; le traitement IHM recherche aussi les références dans les snippets."),
            ("Traiter les variables IHM", "Non", "Génère les fichiers .flvar depuis le XML Author-it sélectionné et remplace les références de snippets reconnues par des références MadCap:variable."),
            ("Nettoyer la TOC sans modifier sa hiérarchie", "Oui", "Supprime les entrées parasites du profil, par exemple A_HEADER, A_FOOTER, Table des matières ou Cover. Les niveaux utiles restent imbriqués."),
            ("Vérifier la target sans la modifier", "Oui", "Contrôle MasterToc, MasterStylesheet, MasterPageLayout et, pour PDF/Word, PatchHeadingLevels. Aucun fichier n'est écrit."),
            ("Réparer TOC, CSS, layout et niveaux de titres", "Non", "Si la vérification trouve un écart, corrige uniquement ces réglages. Les conditions et variables de target sont conservées."),
            ("Générer le rapport final unifié", "Oui", "Crée un journal récapitulant les options, transformations, suppressions TOC, écarts de target, avertissements et erreurs."),
        ],
        [3200, 900, 5260],
    )
    add_callout(doc, "RÈGLE DE SÉCURITÉ", "Laisser Réparer décoché lors du premier passage. Lire le rapport de vérification, puis relancer avec Réparer uniquement si la target présente un écart confirmé.", fill="FFF7E6", accent=GOLD)

    add_heading(doc, "4.3 Options Cleanup", 2)
    add_table(
        doc,
        ["Option", "Transformation effectuée"],
        [
            ("Actions / résultats", "Reconstruit les listes d'actions numérotées ou à puces et rattache les résultats correspondants."),
            ("Listes à tirets", "Convertit les paragraphes importés en vraies listes, avec leurs niveaux d'imbrication."),
            ("Callouts", "Transforme les encadrés importés en structures ATLAS et utilise les pictogrammes livrés dans les ressources."),
            ("Figures", "Regroupe les images et leurs légendes dans les blocs figure attendus."),
            ("Styles simples", "Normalise les classes et styles Author-it simples, notamment centrage, indice et exposant."),
        ],
        [2500, 6860],
    )

    add_heading(doc, "4.4 Source IHM", 2)
    add_body(doc, "Cette zone devient active uniquement lorsque Traiter les variables IHM est coché.")
    for text in (
        "Cliquer sur XML... et sélectionner le fichier XML original exporté depuis Author-it, pas un topic .htm de Flare.",
        "Attendre la détection des templates IHM français.",
        "Cocher uniquement les templates à convertir.",
        "Après exécution, vérifier les fichiers .flvar générés et les compteurs de références remplacées ou non associées.",
    ):
        add_list_item(doc, text, number_id)
    add_callout(doc, "CONDITIONS IHM", "La détection des variantes conditionnelles internes au XML n'est pas considérée comme stabilisée. Ne pas activer ce traitement sans XML validé et consigne métier claire.", fill="FDECEC", accent=RED)

    add_heading(doc, "4.5 Procédure recommandée sans IHM", 2)
    for text in (
        "Ouvrir un topic du projet importé, puis cliquer sur Finaliser import AIT.",
        "Choisir le bon type documentaire.",
        "Sélectionner la target du document et vérifier que la TOC automatiquement proposée est la bonne.",
        "Limiter Dossier à nettoyer au dossier importé si le reste de Content ne doit pas être scanné.",
        "Conserver les options cochées par défaut et laisser Traiter les variables IHM et Réparer décochés.",
        "Cliquer sur Finaliser et lire le récapitulatif.",
        "Si la target est déclarée non conforme, vérifier le détail du rapport puis relancer avec Réparer coché.",
        "Ouvrir la TOC et la target dans Flare, puis construire le PDF.",
    ):
        add_list_item(doc, text, number_id)

    add_heading(doc, "4.6 Cas particulier : TOC plate dans le PDF", 2)
    add_body(doc, "Si la TOC Flare est correctement imbriquée mais que le PDF numérote tous les topics au même niveau (2, 3, 4...) au lieu de 2.1, 2.2 et 2.2.1, la target imprimée ne possède probablement pas PatchHeadingLevels=\"true\".")
    add_callout(doc, "CORRECTION", "Sélectionner la target PDF, cocher Vérifier puis Réparer, finaliser et reconstruire le PDF. Le plug-in ajoute PatchHeadingLevels uniquement aux targets PDF/Word.")

    add_heading(doc, "4.7 Journaux et sauvegardes", 2)
    add_table(
        doc,
        ["Élément", "Emplacement / suffixe"],
        [
            ("Rapport final unifié", "<Projet>/Project/AITWorkflowLogs/AIT_Finalization_yyyyMMdd_HHmmss.log"),
            ("Journal Cleanup détaillé", "Documents/ATLASDocGenerator/AITCleanupLogs/AIT_Cleanup_Log_yyyyMMdd_HHmmss.txt"),
            ("Sauvegarde TOC ou target", "Fichier d'origine suivi de .bak"),
            ("Sauvegarde d'une ressource remplacée", "Fichier d'origine suivi de .before-ait-finalizer.bak"),
            ("General.flvar", "Toujours préservé lorsqu'il existe déjà dans le projet."),
        ],
        [2900, 6460],
    )
    add_placeholder(
        doc,
        "Figure 5 — Récapitulatif de fin de Finaliser import AIT",
        "Après un essai réussi dans Flare, capturer la boîte de dialogue indiquant ressources, transformations, TOC, target, avertissements, erreurs et chemin du rapport.",
    )

    # Troubleshooting
    add_heading(doc, "5. Dépannage rapide", 1, page_break=True)
    add_table(
        doc,
        ["Symptôme", "Cause probable", "Action"],
        [
            ("Le bouton refuse de s'ouvrir", "Aucun topic du projet n'est actif.", "Ouvrir un topic appartenant au projet Flare courant."),
            ("Aucune target dans Checklist Generator", "Les targets n'ont pas de MasterToc valide.", "Ouvrir la target et corriger sa TOC principale."),
            ("Aucun H1 admissible", "La TOC ne référence pas les bons topics, les H1 sont no_num ou les prérequis n'utilisent pas p.ss_section.", "Corriger la structure source puis relancer."),
            ("QIQO_table introuvable", "Le snippet requis manque dans QIQO_content.", "Installer les ressources ATLAS ou restaurer le snippet depuis le projet parent."),
            ("Topic modèle introuvable", "Les ressources communes du projet sont incomplètes.", "Installer les ressources et vérifier Content/Resources/Commun Stago/Topics_Tech."),
            ("Callout sans pictogramme", "Images ou CSS ATLAS absentes/non liées.", "Installer les ressources, vérifier Styles.css et les chemins d'images."),
            ("Target non conforme", "TOC, CSS, layout ou niveaux de titres incorrects.", "Lire le rapport ; cocher Réparer seulement après confirmation."),
            ("TOC PDF sans niveaux", "PatchHeadingLevels absent de la target PDF.", "Vérifier et réparer la target, puis reconstruire le PDF."),
            ("Aucun template IHM détecté", "Le mauvais fichier a été sélectionné ou le XML ne contient pas de template français reconnu.", "Choisir l'export XML Author-it original et contrôler son contenu."),
            ("Un fichier existe déjà", "Le nom du package ou de la checklist est déjà utilisé.", "Choisir une autre référence ou traiter explicitement l'ancien fichier."),
        ],
        [2300, 3420, 3640],
    )

    add_heading(doc, "6. Checklist de validation finale", 1)
    for text in (
        "Le bon projet était actif au lancement.",
        "La TOC sélectionnée contient les bons topics et conserve ses niveaux.",
        "La target sélectionnée pointe vers cette TOC.",
        "Les conditions et variables de target attendues sont toujours présentes.",
        "General.flvar n'a pas été remplacé.",
        "Les liens, snippets, images, callouts et figures s'affichent dans Flare.",
        "Le journal ne contient aucune erreur non expliquée.",
        "Le PDF de contrôle est conforme, y compris la numérotation des titres.",
    ):
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "7. Emplacements des captures complémentaires", 1)
    add_body(doc, "Les figures 2, 3 et 4 sont déjà intégrées au guide à partir des formulaires réels du plug-in. Deux captures pourront être ajoutées après test dans Flare :")
    add_list_item(doc, "Figure 1, section 1.1 : l'onglet ATLAS complet dans le ruban Flare.", bullet_id)
    add_list_item(doc, "Figure 5, section 4.7 : le récapitulatif final après un import AIT réussi.", bullet_id)

    doc.core_properties.title = "Guide d'utilisation ATLAS Doc Generator"
    doc.core_properties.subject = "Utilisation des trois boutons du plug-in MadCap Flare"
    doc.core_properties.author = "Documentation ATLAS"
    doc.core_properties.keywords = "ATLAS, MadCap Flare, Doc Generator, Checklist, Author-it"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
